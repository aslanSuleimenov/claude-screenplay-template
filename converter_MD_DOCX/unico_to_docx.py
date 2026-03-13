"""
Converter analytics/unico_package.md → versions/unico_vNN.docx

Usage:
    python converter_MD_DOCX/unico_to_docx.py
    python converter_MD_DOCX/unico_to_docx.py MyProject_unico_v2

Requires: python-docx
    pip install python-docx
"""

import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Design constants ──────────────────────────────────────────────────────────

FONT_MAIN  = "Arial"
FONT_TITLE = "Arial"

PT_COVER_LABEL   = 10
PT_COVER_TITLE   = 28
PT_COVER_SUB     = 12
PT_SECTION       = 13
PT_SUBSECTION    = 11
PT_FIELD_LABEL   = 10
PT_BODY          = 10
PT_CHARACTER_NAME = 12

COLOR_BRAND    = RGBColor(0x1a, 0x1a, 0x2e)   # dark navy — UNICO header color
COLOR_ACCENT   = RGBColor(0xc8, 0x00, 0x28)   # crimson — UNICO accent
COLOR_LABEL    = RGBColor(0x44, 0x44, 0x44)   # dark grey for field labels
COLOR_BODY     = RGBColor(0x1a, 0x1a, 0x1a)   # near-black body text
COLOR_RULE     = RGBColor(0xcc, 0xcc, 0xcc)   # light grey rule


# ── Helpers ───────────────────────────────────────────────────────────────────

def add_rule(doc, color_hex="CCCCCC", space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name   = FONT_MAIN
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_para(doc, text, size_pt, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
             space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size_pt, bold, italic, color)
    return p


def parse_inline(paragraph, text, size_pt, base_color=None, base_italic=False):
    """Handle **bold** and *italic* inline markers."""
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size_pt, bold=True, italic=base_italic, color=base_color)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size_pt, italic=True, color=base_color)
        elif token:
            run = paragraph.add_run(token)
            set_font(run, size_pt, italic=base_italic, color=base_color)


def set_page(doc):
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)


# ── Cover page ────────────────────────────────────────────────────────────────

def add_cover(doc, title, author, date):
    set_page(doc)

    # Remove default empty paragraph
    for el in list(doc.element.body):
        doc.element.body.remove(el)

    # Top accent rule
    add_rule(doc, color_hex="C80028", space_before=0, space_after=0)

    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Label
    add_para(doc, "UNICO", PT_COVER_LABEL, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_ACCENT,
             space_before=0, space_after=4)

    # Project title
    add_para(doc, title, PT_COVER_TITLE, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_BRAND,
             space_before=0, space_after=6)

    add_rule(doc, color_hex="C80028", space_before=4, space_after=4)

    # Subtitle
    add_para(doc, "PITCHING STARTER PACK", PT_COVER_SUB,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_LABEL,
             space_before=0, space_after=30)

    # Author / date
    add_para(doc, f"Author: {author}", PT_BODY,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_LABEL)
    add_para(doc, f"Date: {date}", PT_BODY,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_LABEL)

    # Page break
    doc.add_page_break()


# ── Main converter ────────────────────────────────────────────────────────────

def convert_unico(source_path, output_path):
    text  = source_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    set_page(doc)
    for el in list(doc.element.body):
        doc.element.body.remove(el)

    # Extract title/author/date from first lines for cover
    title  = "Untitled"
    author = "—"
    date   = "—"
    for line in lines[:10]:
        if line.startswith("# UNICO Starter Pack:"):
            title = line.split(":", 1)[1].strip()
        m = re.match(r'\*?Author:\*?\s*(.*)', line, re.IGNORECASE)
        if m:
            author = m.group(1).strip()
        m = re.match(r'\*?Date:\*?\s*(.*)', line, re.IGNORECASE)
        if m:
            date = m.group(1).strip()

    add_cover(doc, title, author, date)

    i = 0
    in_character = False
    character_name = ""

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip cover metadata lines (already used)
        if (line.startswith("# UNICO Starter Pack:") or
                re.match(r'(Author|Date):\s', line) or
                line == "---"):
            if line == "---":
                add_rule(doc, space_before=8, space_after=8)
            i += 1
            continue

        # Empty line
        if line == "":
            i += 1
            continue

        # H2: section header (PROJECT PASSPORT / CHARACTER BIBLE / etc.)
        if line.startswith("## "):
            section = line[3:].strip().upper()
            in_character = (section == "CHARACTER BIBLE")
            add_para(doc, section, PT_SECTION, bold=True,
                     color=COLOR_BRAND, space_before=16, space_after=6)
            add_rule(doc, color_hex="C80028", space_before=0, space_after=8)
            i += 1
            continue

        # H3: character name inside Character Bible
        if line.startswith("### "):
            character_name = line[4:].strip()
            in_character = True
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(character_name.upper())
            set_font(run, PT_CHARACTER_NAME, bold=True, color=COLOR_ACCENT)
            add_rule(doc, color_hex="CCCCCC", space_before=0, space_after=6)
            i += 1
            continue

        # Numbered list (presentation slides)
        num_match = re.match(r'^(\d+)\.\s+\*\*([^*]+)\*\*\s*[—–-]?\s*(.*)', line)
        if num_match:
            num, label, rest = num_match.groups()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.5)
            r1 = p.add_run(f"{num}. ")
            set_font(r1, PT_BODY, bold=True, color=COLOR_ACCENT)
            r2 = p.add_run(label)
            set_font(r2, PT_BODY, bold=True, color=COLOR_LABEL)
            if rest:
                r3 = p.add_run(f" — {rest}")
                set_font(r3, PT_BODY, color=COLOR_BODY)
            i += 1
            continue

        # Bullet list
        if line.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.5)
            run = p.add_run("• ")
            set_font(run, PT_BODY, color=COLOR_ACCENT)
            parse_inline(p, line[2:], PT_BODY, base_color=COLOR_BODY)
            i += 1
            continue

        # Field: **Label:** value
        field_match = re.match(r'^\*\*([^*]+):\*\*\s*(.*)', line)
        if field_match:
            label, value = field_match.group(1), field_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            r1 = p.add_run(label + ": ")
            set_font(r1, PT_FIELD_LABEL, bold=True, color=COLOR_LABEL)
            if value:
                parse_inline(p, value, PT_BODY, base_color=COLOR_BODY)
            i += 1
            # Collect continuation lines (indented or plain text until next field/section)
            while i < len(lines):
                cont = lines[i].rstrip()
                if (cont == "" or cont.startswith("#") or
                        re.match(r'^\*\*[^*]+:\*\*', cont) or
                        cont.startswith("- ") or cont.startswith("---")):
                    break
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after  = Pt(2)
                p2.paragraph_format.left_indent  = Cm(0.3)
                parse_inline(p2, cont, PT_BODY, base_color=COLOR_BODY)
                i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        parse_inline(p, line, PT_BODY, base_color=COLOR_BODY)
        i += 1

    doc.save(str(output_path))
    print(f"Saved: {output_path}")


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    root   = Path(__file__).parent.parent
    source = root / "analytics" / "unico_package.md"

    if not source.exists():
        print(f"Error: {source} not found.")
        print("Run the unico agent first to generate unico_package.md")
        sys.exit(1)

    versions_dir = root / "versions"
    versions_dir.mkdir(exist_ok=True)

    if len(sys.argv) > 1:
        stem = sys.argv[1]
    else:
        project_name = root.name
        existing = list(versions_dir.glob(f"{project_name}_unico_v*.docx"))
        nums = [int(m.group(1)) for f in existing
                if (m := re.search(r'_unico_v(\d+)\.docx$', f.name))]
        next_v = (max(nums) + 1) if nums else 1
        stem = f"{project_name}_unico_v{next_v:02d}"

    output = versions_dir / f"{stem}.docx"
    convert_unico(source, output)


if __name__ == "__main__":
    main()
