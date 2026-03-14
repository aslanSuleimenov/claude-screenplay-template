"""Конвертер scenes/*.md → DOCX.

Автоматически определяет тип проекта из CLAUDE.md:
  - художественный → однокол​оночный формат (screenplay)
  - документальный → двухкол​оночный формат (AV-script)

Использование:
  python converter_MD_DOCX/md_to_docx.py [имя_файла]
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

FONT = "Courier New"
SZ = Pt(12)
SZ_SCENE = Pt(12)
SZ_SUB = Pt(14)
SZ_TITLE = Pt(24)

BASE = Path(__file__).parent.parent
# If launched from the project root (e.g. via `python converter_MD_DOCX/md_to_docx.py`),
# prefer cwd so relative paths work regardless of where the script physically lives.
if (Path.cwd() / "scenes").exists():
    BASE = Path.cwd()
scenes_dir = BASE / "scenes"
versions_dir = BASE / "versions"


# --- Определить тип проекта ---
def detect_project_type():
    claude_md = BASE / "CLAUDE.md"
    if not claude_md.exists():
        return "художественный"
    text = claude_md.read_text(encoding="utf-8-sig")
    m = re.search(r'\*\*Тип:\*\*\s*(\S+)', text)
    if m and "документ" in m.group(1).lower():
        return "документальный"
    return "художественный"


def find_next_version(proj_name):
    """Найти следующий номер версии, безопасно обрабатывая нечисловые суффиксы."""
    existing = sorted(versions_dir.glob(f"{proj_name}_v*.docx"))
    max_num = 0
    for f in existing:
        suffix = f.stem.rsplit("_v", 1)[-1]
        try:
            num = int(suffix)
            if num > max_num:
                max_num = num
        except ValueError:
            continue  # пропустить файлы с нечисловым суффиксом
    return max_num + 1


def add_page_numbers(doc):
    """Добавить номера страниц в footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # PAGE field
        run = p.add_run()
        fld_char_begin = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run._r.append(fld_char_begin)
        run2 = p.add_run()
        instr = run2._r.makeelement(qn('w:instrText'), {})
        instr.text = ' PAGE '
        run2._r.append(instr)
        run3 = p.add_run()
        fld_char_end = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run3._r.append(fld_char_end)
        for r in (run, run2, run3):
            r.font.name = FONT
            r.font.size = Pt(10)


# ============================================================
#  ОБЩИЕ ФУНКЦИИ
# ============================================================

def add_run(p, text, bold=None, italic=None, size=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = size or SZ
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic
    return r


def setup_page(doc, orientation="portrait"):
    if orientation == "landscape":
        page_w = Emu(10692130)
        page_h = Emu(7560310)
        margin_left = Inches(1.0)
        margin_right = Inches(1.0)
    else:
        page_w = Emu(7560310)
        page_h = Emu(10692130)
        margin_left = Inches(1.5)
        margin_right = Inches(1.0)

    for sec in doc.sections:
        sec.page_width = page_w
        sec.page_height = page_h
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = margin_left
        sec.right_margin = margin_right
        if orientation == "landscape":
            sec.orientation = WD_ORIENT.LANDSCAPE


def make_title_page(doc):
    """Генерация титульной страницы из CLAUDE.md."""
    claude_md = BASE / "CLAUDE.md"
    title = "МОИ ИГРУШКИ"
    logline = ""
    genre = ""
    if claude_md.exists():
        text = claude_md.read_text(encoding="utf-8-sig")
        m = re.search(r'\*\*Название:\*\*\s*(.+)', text)
        if m:
            title = m.group(1).strip()
        m = re.search(r'\*\*Логлайн:\*\*\s*(.+)', text)
        if m:
            logline = m.group(1).strip()
        m = re.search(r'\*\*Жанр:\*\*\s*(.+)', text)
        if m:
            genre = m.group(1).strip()

    # Отступ сверху ~1/3 страницы
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(180)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    add_run(p, title.upper(), bold=True, size=SZ_TITLE)

    if genre:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        add_run(p, genre, size=SZ)

    if logline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        add_run(p, logline, italic=True, size=SZ)

    doc.add_page_break()


# ============================================================
#  ХУДОЖЕСТВЕННЫЙ ФОРМАТ (screenplay)
# ============================================================

INDENT_CHAR = Inches(2.0)
INDENT_PAREN = Inches(1.5)
INDENT_DIAL = Inches(1.0)


def is_character_line(s):
    if s.startswith("**") and s.endswith("**") and len(s) < 80:
        inner = s.replace("**", "").strip()
        if any(kw in inner.upper() for kw in ["ИНТ.", "НАТ.", "INT.", "EXT.", "МОНТАЖ", "НАДПИСЬ"]):
            return False
        return bool(inner)
    clean = re.sub(r'\s*\(.*\)\s*$', '', s.replace("**", "").strip())
    if (clean.isupper() and 2 < len(clean) < 40
            and "." not in clean and clean not in ("КОНЕЦ", "КОНЕЦ.")):
        return True
    return False


def extract_character_name(s):
    clean = s.replace("**", "").strip()
    m = re.match(r'^(.+?)\s*\((.+)\)\s*$', clean)
    if m:
        return m.group(1).strip().upper(), m.group(2).strip()
    return clean.upper(), None


def convert_fiction(doc, scene_files):
    setup_page(doc, "portrait")
    make_title_page(doc)
    scene_count = 0
    prev_was_blank = False
    last_type = None

    for filepath in scene_files:
        file_lines = filepath.read_text(encoding="utf-8-sig").splitlines()

        for line in file_lines:
            s = line.strip()

            if (s.startswith("<div") or s.startswith("<!--")
                    or s == "---" or s == "===SCENE_BREAK==="):
                continue

            if not s:
                if not prev_was_blank:
                    prev_was_blank = True
                continue

            had_blank = prev_was_blank
            prev_was_blank = False
            # Пустая строка разрывает диалог — сбросить тип
            if had_blank and last_type in ("dialogue", "paren"):
                last_type = "action"

            clean = s.replace("**", "").replace("*", "").strip()

            if clean.startswith("ЗАТЕМНЕНИЕ") or clean == "CUT TO:":
                continue

            # Заголовок сцены
            if s.startswith("# "):
                text = s[2:].strip()
                scene_count += 1
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(24)
                display = text.upper()
                if not display.startswith("СЦЕНА"):
                    display = f"СЦЕНА {scene_count}: {display}"
                add_run(p, display, bold=True, size=SZ_SCENE)
                p.paragraph_format.space_after = Pt(6)
                last_type = "scene_head"
                continue

            # H2
            if s.startswith("## "):
                text = s[3:].strip().replace("**", "").replace("*", "")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                add_run(p, text.upper(), bold=True, size=SZ_SUB)
                last_type = "action"
                continue

            # Slug line
            if any(kw in s.upper() for kw in ["ИНТ.", "НАТ.", "INT.", "EXT."]) or \
                    (s.startswith("**") and any(kw in s.upper() for kw in ["МОНТАЖ", "НАДПИСЬ"])):
                text = s.replace("**", "").strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                add_run(p, text.upper(), bold=True)
                last_type = "slug"
                continue

            # Ремарка
            if s.startswith("*(") and s.endswith(")*"):
                inner = s[1:-1].strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = INDENT_PAREN
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                add_run(p, inner, italic=True)
                last_type = "paren"
                continue

            # Курсив
            if s.startswith("*") and s.endswith("*") and not s.startswith("**"):
                inner = s[1:-1].strip()
                if inner.startswith("(") and inner.endswith(")"):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = INDENT_PAREN
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    add_run(p, inner, italic=True)
                    last_type = "paren"
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    add_run(p, inner, italic=True)
                    last_type = "action"
                continue

            # Имя персонажа
            if is_character_line(s):
                name, paren = extract_character_name(s)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = INDENT_CHAR
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(0)
                add_run(p, f"{name} ({paren})" if paren else name, bold=True)
                last_type = "character"
                continue

            # Блок цитаты (заметки, палитра) — пропускаем
            if s.startswith("> "):
                continue

            # Обычный текст
            text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', s)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)

            p = doc.add_paragraph()
            if last_type in ("character", "paren", "dialogue"):
                p.paragraph_format.left_indent = INDENT_DIAL
                p.paragraph_format.right_indent = INDENT_DIAL
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                add_run(p, text)
                last_type = "dialogue"
            else:
                p.paragraph_format.space_before = Pt(12) if had_blank else Pt(0)
                p.paragraph_format.space_after = Pt(0)
                add_run(p, text)
                last_type = "action"

    return scene_count


# ============================================================
#  ДОКУМЕНТАЛЬНЫЙ ФОРМАТ (AV-script)
# ============================================================

def set_cell_font(cell, text, bold=False, italic=False, size=None):
    """Установить текст и шрифт в ячейке таблицы."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = size or SZ
    r.bold = bold
    r.italic = italic
    return r


def set_cell_rich(cell, text, size=None):
    """Парсит markdown-разметку (**bold**, *italic*) в ячейке."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

    # Разбиваем по **bold** и *italic* маркерам
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1])
            r.italic = True
        else:
            r = p.add_run(part)
        r.font.name = FONT
        r.font.size = size or SZ


def parse_md_table(lines, start_idx):
    """Парсит markdown-таблицу начиная с заголовка | VIDEO | AUDIO |.
    Возвращает (rows, end_idx). rows = [(video_text, audio_text), ...]"""
    rows = []
    i = start_idx
    # Пропустить заголовок и разделитель
    if i < len(lines) and "|" in lines[i]:
        i += 1  # | VIDEO | AUDIO |
    if i < len(lines) and re.match(r'\|[-\s|]+\|', lines[i]):
        i += 1  # |-------|-------|

    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        cells = line.strip("|").split("|")
        cells = [c.strip() for c in cells]

        video = cells[0] if len(cells) > 0 else ""
        audio = cells[1] if len(cells) > 1 else ""
        rows.append((video, audio))
        i += 1

    return rows, i


def convert_documentary(doc, scene_files):
    setup_page(doc, "landscape")
    block_count = 0

    for filepath in scene_files:
        file_lines = filepath.read_text(encoding="utf-8-sig").splitlines()

        i = 0
        while i < len(file_lines):
            s = file_lines[i].strip()

            # Пропуск
            if not s or s == "---" or s.startswith("<!--") or s.startswith("<div"):
                i += 1
                continue

            # Заголовок блока
            if s.startswith("# "):
                text = s[2:].strip()
                block_count += 1
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
                display = text.upper()
                if not display.startswith("БЛОК"):
                    display = f"БЛОК {block_count}: {display}"
                add_run(p, display, bold=True, size=SZ_SCENE)
                i += 1
                continue

            # H2: титр-разделитель
            if s.startswith("## "):
                text = s[3:].strip().replace("**", "").replace("*", "")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                add_run(p, text.upper(), bold=True, size=SZ_SUB)
                i += 1
                continue

            # Режиссёрская заметка — пропускаем
            if s.startswith("> "):
                i += 1
                continue

            # Таблица VIDEO | AUDIO
            if "|" in s and re.search(r'VIDEO|AUDIO|video|audio', s, re.IGNORECASE):
                rows, end_idx = parse_md_table(file_lines, i)
                if rows:
                    table = doc.add_table(rows=len(rows), cols=2)
                    table.style = "Table Grid"

                    # Ширина колонок: VIDEO 55%, AUDIO 45%
                    sec = doc.sections[-1]
                    avail_width = sec.page_width - sec.left_margin - sec.right_margin
                    video_w = int(avail_width * 0.55)
                    audio_w = int(avail_width * 0.45)

                    for row_idx, (video, audio) in enumerate(rows):
                        video_cell = table.cell(row_idx, 0)
                        audio_cell = table.cell(row_idx, 1)
                        video_cell.width = video_w
                        audio_cell.width = audio_w
                        set_cell_rich(video_cell, video)
                        set_cell_rich(audio_cell, audio)

                    # Отступ после таблицы
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)

                i = end_idx
                continue

            # Обычный текст (авторские заметки между таблицами)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', s)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            add_run(p, text, italic=True, size=Pt(10))
            i += 1

    return block_count


# ============================================================
#  MAIN
# ============================================================

def main():
    versions_dir.mkdir(exist_ok=True)

    proj_name = sys.argv[1] if len(sys.argv) > 1 else BASE.name
    project_type = detect_project_type()

    scene_files = sorted(scenes_dir.glob("*.md"))
    if not scene_files:
        print("Нет .md файлов в scenes/")
        sys.exit(1)

    next_num = find_next_version(proj_name)
    dst = versions_dir / f"{proj_name}_v{next_num:02d}.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = SZ
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    print(f"Тип проекта: {project_type}")

    if project_type == "документальный":
        count = convert_documentary(doc, scene_files)
        label = "Блоков"
    else:
        count = convert_fiction(doc, scene_files)
        label = "Сцен"

    add_page_numbers(doc)

    doc.save(str(dst))
    print(f"Сохранено: {dst}")
    print(f"Файлов: {len(scene_files)}")
    print(f"{label}: {count}")
    print(f"Параграфов: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
